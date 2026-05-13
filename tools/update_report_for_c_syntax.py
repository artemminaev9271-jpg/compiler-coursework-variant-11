from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Отчёт_мини_компилятор_вариант_11_проверенный.docx"
INTERMEDIATE = ROOT / "Отчёт_компилятор_вариант_11.docx"
OUT = ROOT / "Отчёт_компилятор_вариант_11_проверенный.docx"
DIAGRAM = ROOT / "diagram_grammar_variant_11.png"


REPLACEMENTS = {
    "мини-компилятор": "компилятор",
    "Мини-компилятор": "Компилятор",
    "мини _ компилятор": "компилятор",
    "Разработка мини-компилятора": "Разработка компилятора",
    "fn compiled_fn(int arg) -> int": "int compiled_fn(int arg)",
    "fn name(int arg) -> int": "int name(int arg)",
    "fn IDENT '(' param_list ')' -> type block": "type IDENT '(' param_list ')' block",
    "function -> fn IDENT": "function -> type IDENT",
}


def replace_text(text: str) -> str:
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def replace_in_paragraph(paragraph) -> None:
    old = paragraph.text
    new = replace_text(old)
    if new != old:
        paragraph.text = new


def replace_in_docx_text() -> None:
    doc = Document(SRC)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    doc.save(INTERMEDIATE)


def replace_diagram_media() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with ZipFile(INTERMEDIATE, "r") as zin:
            zin.extractall(tmp_path)

        media = tmp_path / "word" / "media" / "image1.png"
        if media.exists():
            shutil.copyfile(DIAGRAM, media)

        with ZipFile(OUT, "w", ZIP_DEFLATED) as zout:
            for path in tmp_path.rglob("*"):
                if path.is_file():
                    zout.write(path, path.relative_to(tmp_path).as_posix())


if __name__ == "__main__":
    replace_in_docx_text()
    replace_diagram_media()
    print(OUT)
