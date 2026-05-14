from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Отчёт_компилятор_вариант_11_проверенный.docx"
REPO_URL = "https://github.com/artemminaev9271-jpg/compiler-coursework-variant-11"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def set_body_style(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def set_toc_style(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def insert_after(paragraph, text: str):
    new_p = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_p._p)
    new_p.text = text
    return new_p


def find_heading(doc: Document, title: str):
    for paragraph in doc.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        if paragraph.text.strip() == title and style.startswith("Heading"):
            return paragraph
    raise RuntimeError(f"Heading not found: {title}")


def replace_text(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            set_body_style(paragraph)


def fix_toc(doc: Document) -> None:
    toc_heading = find_heading(doc, "СОДЕРЖАНИЕ")
    first_heading = find_heading(doc, "1. Задание")

    paragraphs = list(doc.paragraphs)
    start_idx = next(i for i, p in enumerate(paragraphs) if p._p is toc_heading._p)
    first_idx = next(i for i, p in enumerate(paragraphs) if p._p is first_heading._p)

    for paragraph in paragraphs[start_idx + 1:first_idx]:
        remove_paragraph(paragraph)

    toc_items = [
        "1. Задание",
        "2. Лексическая спецификация",
        "3. Диаграмма и описание правил грамматики",
        "4. Реализация лексического и синтаксического анализаторов",
        "5. Описание семантических правил языка",
        "6. Реализация семантического анализатора",
        "7. Описание процесса перевода абстрактного синтаксического дерева (AST) в промежуточное представление LLVM (LLVM IR) и далее в ассемблерный и объектный файлы",
        "8. Реализация генератора кода",
        "9. Тестовые наборы программ",
        "10. Описание процесса использования компилятора",
        "11. Результаты тестирования",
        "12. Заключение",
    ]

    cursor = toc_heading
    for item in toc_items:
        cursor = insert_after(cursor, item)
        set_toc_style(cursor)


def remove_appendices(doc: Document) -> None:
    try:
        appendix_heading = find_heading(doc, "13. Приложения")
    except RuntimeError:
        return

    paragraphs = list(doc.paragraphs)
    start_idx = next(i for i, p in enumerate(paragraphs) if p._p is appendix_heading._p)
    for paragraph in paragraphs[start_idx:]:
        remove_paragraph(paragraph)


def add_repo_note(doc: Document) -> None:
    section4 = find_heading(doc, "4. Реализация лексического и синтаксического анализаторов")
    section5 = find_heading(doc, "5. Описание семантических правил языка")
    paragraphs = list(doc.paragraphs)
    idx4 = next(i for i, p in enumerate(paragraphs) if p._p is section4._p)
    idx5 = next(i for i, p in enumerate(paragraphs) if p._p is section5._p)

    repo_note = (
        "Полные листинги исходных файлов в текст отчёта не включаются. "
        f"Они представлены в репозитории GitHub: {REPO_URL}"
    )

    existing = None
    for paragraph in paragraphs[idx4 + 1:idx5]:
        if REPO_URL in paragraph.text:
            existing = paragraph
            break

    if existing is None:
        target = paragraphs[idx5 - 1]
        new_p = insert_after(target, repo_note)
        set_body_style(new_p)
    else:
        existing.text = repo_note
        set_body_style(existing)


def main() -> None:
    doc = Document(REPORT)
    fix_toc(doc)
    remove_appendices(doc)
    replace_text(
        doc,
        "Ниже приведены ключевые фрагменты Flex и Bison. Полные файлы вынесены в приложения.",
        "Ниже приведены ключевые фрагменты Flex и Bison. Полные листинги исходных файлов представлены в репозитории GitHub.",
    )
    replace_text(
        doc,
        "В листинге показаны фрагменты генерации LLVM IR и объектного файла. Полный файл находится в приложении.",
        "В листинге показаны фрагменты генерации LLVM IR и объектного файла. Полный исходный файл представлен в репозитории GitHub.",
    )
    add_repo_note(doc)
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
