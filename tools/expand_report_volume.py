from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Отчёт_компилятор_вариант_11_проверенный.docx"
GUIDE = ROOT / "Инструкция_скриншоты.md"
WORKTREE_WSL = '/mnt/c/Users/minae/Documents/compiler-coursework-variant-11'


def set_times_new_roman(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_times_new_roman(run, 12)


def format_heading(paragraph, level: int) -> None:
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_times_new_roman(run, 14 if level == 1 else 12, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_run_text(paragraph, text: str, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    set_times_new_roman(run, size, bold=bold, italic=italic)


def insert_after(paragraph, text: str = "", *, style: str | None = None, heading_level: int | None = None):
    new_p = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_p._p)
    if text:
        add_run_text(new_p, text)
    if heading_level is not None:
        format_heading(new_p, heading_level)
    elif style:
        new_p.style = style
    else:
        format_body(new_p)
    return new_p


def add_code_after(paragraph, text: str):
    code_p = insert_after(paragraph, "")
    code_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_p.paragraph_format.first_line_indent = Cm(0)
    code_p.paragraph_format.left_indent = Cm(0.5)
    code_p.paragraph_format.space_after = Pt(6)
    code_p.text = ""
    run = code_p.add_run(text.rstrip())
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)
    return code_p


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_times_new_roman(run, 11, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table_after(paragraph, headers: list[str], rows: list[list[str]]):
    doc = paragraph._parent
    table = doc.add_table(rows=1, cols=len(headers), width=Cm(16))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True)
        set_cell_fill(hdr_cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    paragraph._p.addnext(table._tbl)
    spacer = paragraph.insert_paragraph_before("")
    table._tbl.addnext(spacer._p)
    format_body(spacer)
    return spacer


def add_placeholder_after(paragraph, title: str, command: str, expected: str):
    doc = paragraph._parent
    table = doc.add_table(rows=3, cols=1, width=Cm(16))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        f"МЕСТО ДЛЯ СКРИНШОТА: {title}",
        f"Что должно быть видно: {expected}",
        f"Команда/действие: {command}",
    ]
    for idx, row_text in enumerate(rows):
        cell = table.cell(idx, 0)
        set_cell_text(cell, row_text, bold=idx == 0)
        set_cell_fill(cell, "EAF2F8")
    paragraph._p.addnext(table._tbl)
    spacer = paragraph.insert_paragraph_before("")
    table._tbl.addnext(spacer._p)
    format_body(spacer)
    return spacer


def find_heading(doc: Document, title: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == title:
            return paragraph
    raise RuntimeError(f"Heading not found: {title}")


def has_heading(doc: Document, title: str) -> bool:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if paragraph.text.strip() == title and style_name.startswith("Heading"):
            return True
    return False


def last_paragraph_before(doc: Document, next_heading_title: str):
    next_heading = find_heading(doc, next_heading_title)
    paragraphs = list(doc.paragraphs)
    idx = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is next_heading._p)
    return paragraphs[idx - 1]


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def add_extended_sections(doc: Document) -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "13. Приложения" in full_text and "Подсистема разбора и построения AST" in full_text:
        return

    # Section 1 additions.
    cursor = last_paragraph_before(doc, "2. Лексическая спецификация")
    cursor = insert_after(cursor, "Архитектура проекта выбрана классической для учебных компиляторов и состоит из пяти последовательных фаз: лексический анализ, синтаксический анализ, построение AST, семантическая проверка и генерация LLVM IR с последующим выпуском объектного файла. Такое разделение позволяет проверять корректность программы до начала генерации машинно-зависимого кода.")
    cursor = insert_after(cursor, "Для варианта 11 ключевым требованием является поддержка тернарного оператора cond ? a : b. В реализованном проекте это требование проходит через все фазы: лексер распознаёт символы '?' и ':', парсер строит отдельный узел TernaryExpr, семантический анализатор проверяет тип условия и совместимость ветвей, а генератор кода создаёт набор basic blocks и PHI-узел в LLVM IR.")
    cursor = insert_after(cursor, "С практической точки зрения проект ориентирован на среду Ubuntu WSL и компоновку с runtime/main.c. Сам компилятор создаёт объектный файл, а запуск демонстрационной программы выполняется уже после линковки с небольшим исполняющим окружением, вызывающим compiled_fn(10).")
    cursor = add_table_after(
        cursor,
        ["Модуль", "Назначение"],
        [
            ["src/lexer.l", "Токенизация входного файла .mc и передача координат токенов в парсер."],
            ["src/parser.y", "Описание грамматики и построение AST без генерации LLVM IR на этапе разбора."],
            ["src/sema.cpp", "Проверка областей видимости, типов, return и вызовов функций."],
            ["src/codegen.cpp", "Преобразование проверенного AST в LLVM IR и объектный файл x86_64."],
            ["src/driver.cpp", "Командная строка, запуск анализа и управление выводом .ll/.o."],
            ["runtime/main.c", "Минимальная оболочка для запуска скомпилированной функции compiled_fn."],
        ],
    )

    # Section 3 additions.
    cursor = last_paragraph_before(doc, "4. Реализация лексического и синтаксического анализаторов")
    cursor = insert_after(cursor, "Грамматика языка построена так, чтобы сохранить привычный C-подобный синтаксис, но при этом оставить реализацию компактной. Поддерживается только тип int, что упрощает семантические проверки и позволяет сосредоточиться на правильности построения управляющего потока, вызовов функций и вычисления выражений.")
    cursor = insert_after(cursor, "Отдельно важно, что правила разбора выражений разбиты по приоритетам. Благодаря этому выражения вида a + b * c, x < y && z != 0 и cond ? a : b интерпретируются однозначно и не требуют дополнительных скобок, если пользователь пишет стандартную форму записи.")
    cursor = add_table_after(
        cursor,
        ["Уровень", "Операторы", "Ассоциативность"],
        [
            ["1", "?:", "справа налево"],
            ["2", "||", "слева направо"],
            ["3", "&&", "слева направо"],
            ["4", "==, !=", "слева направо"],
            ["5", "<, <=, >, >=", "слева направо"],
            ["6", "+, -", "слева направо"],
            ["7", "*, /, %", "слева направо"],
            ["8", "!, унарный -", "справа налево"],
        ],
    )
    cursor = insert_after(cursor, "Такое распределение приоритетов соответствует ожидаемому поведению тестов проекта. В частности, тернарный оператор имеет самый низкий приоритет среди реализованных бинарных и логических операций, поэтому выражения в его условии вычисляются полностью до выбора одной из ветвей.")

    # Section 5 additions.
    cursor = last_paragraph_before(doc, "6. Реализация семантического анализатора")
    cursor = insert_after(cursor, "Семантический анализатор работает поверх уже построенного AST и не зависит от конкретного синтаксиса исходного текста. За счёт этого ошибки формулируются в терминах сущностей языка: переменных, функций, выражений, операторов return и условий.")
    cursor = insert_after(cursor, "В проекте реализован стек областей видимости. При входе в блок создаётся новая таблица символов, а при выходе она удаляется. Это позволяет корректно обрабатывать локальные объявления внутри составных операторов и цикла for, не смешивая их с именами из внешних блоков.")
    cursor = add_table_after(
        cursor,
        ["Группа правил", "Что проверяется"],
        [
            ["Функции", "Уникальность имён, наличие обязательной функции compiled_fn(int), корректность числа параметров."],
            ["Переменные", "Запрет повторного объявления в одной области видимости и запрет обращения к необъявленному имени."],
            ["Условия", "Выражения в if, for, логических операторах и тернарном условии должны иметь тип bool."],
            ["Оператор return", "Тип возвращаемого выражения должен совпадать с типом текущей функции."],
            ["Тернарный оператор", "Обе ветви должны иметь одинаковый тип, который затем присваивается всему выражению."],
        ],
    )
    cursor = insert_after(cursor, "Поскольку язык проекта невелик, именно семантическая фаза отвечает за большую часть пользовательской диагностики. Это видно по отрицательным тестам: ошибки о несоответствии типов и использовании необъявленной переменной выдаются до генерации LLVM IR, что делает поведение компилятора предсказуемым.")

    # Section 7 additions.
    cursor = last_paragraph_before(doc, "8. Реализация генератора кода")
    cursor = insert_after(cursor, "После успешной семантической проверки AST передаётся генератору кода. Его задача состоит не просто в последовательной печати инструкций LLVM IR, а в построении корректного управляющего графа функции, включая ветвления, циклы и слияние потоков управления.")
    cursor = insert_after(cursor, "Для целевой архитектуры выбран target triple x86_64-pc-linux-gnu. Именно под него настраивается LLVM-модуль, после чего объектный файл можно напрямую линковать обычным clang вместе с runtime/main.c без дополнительных прослоек и эмуляторов.")
    cursor = insert_after(cursor, "Особенно показателен путь тернарного оператора. Генератор создаёт блоки then, else и merge, после чего собирает итоговое значение выражения через PHI-узел. Аналогичный принцип используется для логических операторов с коротким замыканием, когда правая часть вычисляется не всегда.")
    cursor = add_placeholder_after(
        cursor,
        "Фрагмент LLVM IR с PHI-узлом",
        f"cd \"{WORKTREE_WSL}/build\"\n./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll\ngrep -n phi generated.ll",
        "Строки с инструкцией phi в generated.ll и имя временного значения ternarytmp либо logictmp.",
    )

    # Section 9 additions.
    cursor = last_paragraph_before(doc, "10. Описание процесса использования компилятора")
    cursor = insert_after(cursor, "Набор тестов проекта подобран так, чтобы покрыть как отдельные конструкции языка, так и интеграцию всех фаз компилятора. В положительных примерах проверяются арифметика, условные операторы, циклы, вызовы функций, короткое замыкание и более содержательная программа вычисления суммы простых чисел.")
    cursor = add_table_after(
        cursor,
        ["Категория", "Файлы", "Цель"],
        [
            ["Базовые", "01_return_literal.mc, 02_arithmetic.mc", "Проверка литералов и базовых арифметических операций."],
            ["Управляющие конструкции", "03_if_else.mc, 04_for.mc", "Проверка ветвлений и циклов."],
            ["Вариант 11", "05_ternary_true.mc, 06_ternary_false.mc, 07_nested_ternary.mc", "Проверка тернарного оператора и его вложенности."],
            ["Диагностика", "08_type_error.mc, 09_undeclared_variable.mc", "Проверка сообщений semantic error."],
            ["Интеграционные", "10_short_circuit.mc, 11_function_call.mc, 12_sum_primes.mc, test.mc", "Проверка кооперации нескольких конструкций языка в одной программе."],
        ],
    )
    cursor = insert_after(cursor, "Тест 12_sum_primes.mc имеет особую ценность для отчёта, поскольку показывает, что компилятор корректно обрабатывает несколько функций, вложенные if и for, арифметические выражения, операции сравнения и возврат накопленного результата. Это уже не минимальный пример, а небольшая законченная программа.")

    # Section 10 additions.
    cursor = last_paragraph_before(doc, "11. Результаты тестирования")
    cursor = insert_after(cursor, "Практический сценарий использования компилятора удобно разбивать на несколько отдельных действий: проверка структуры проекта, сборка, компиляция тестовой программы, просмотр промежуточного LLVM IR и запуск полученного исполняемого файла после линковки. Такой сценарий хорошо подходит и для демонстрации преподавателю.")
    cursor = add_placeholder_after(
        cursor,
        "Структура проекта",
        f"cd \"{WORKTREE_WSL}\"\nls\nfind src runtime tests -maxdepth 1 -type f | sort",
        "Каталоги src, runtime, tests, а также файлы CMakeLists.txt и README.md.",
    )
    cursor = add_placeholder_after(
        cursor,
        "Успешная сборка компилятора",
        f"cd \"{WORKTREE_WSL}\"\nmkdir -p build\ncd build\ncmake ..\ncmake --build .",
        "Сообщения конфигурации CMake и строка об успешной сборке mini_cc.",
    )
    cursor = add_placeholder_after(
        cursor,
        "Запуск основной тестовой программы",
        f"cd \"{WORKTREE_WSL}/build\"\n./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll\nclang ../runtime/main.c generated.o -o app\n./app",
        "Команды компиляции и итоговый вывод 14.",
    )

    # Section 11 additions.
    cursor = last_paragraph_before(doc, "12. Заключение")
    cursor = insert_after(cursor, "Отрицательные тесты показывают, что диагностика реализована не формально, а действительно встроена в компилятор. Ошибка несовместимого инициализатора ловится в момент анализа объявления переменной, а обращение к отсутствующему имени определяется через поиск по стеку областей видимости.")
    cursor = insert_after(cursor, "Положительные интеграционные тесты подтверждают согласованность всех фаз. Например, short-circuit тест одновременно проверяет логический оператор &&, ветвление if/else, корректность генерации условных переходов и отсутствие вычисления правой части выражения при ложном левом операнде.")
    cursor = insert_after(cursor, "Дополнительным сильным аргументом в отчёте служит тест с суммой простых чисел. Он демонстрирует, что генерация кода справляется не только с маленькими выражениями, но и с программой, содержащей две функции, вложенный цикл и проверку делимости через остаток от деления.")
    cursor = add_placeholder_after(
        cursor,
        "Семантические ошибки",
        f"cd \"{WORKTREE_WSL}/build\"\n./mini_cc ../tests/08_type_error.mc -o bad_type.o --emit-ir bad_type.ll\n./mini_cc ../tests/09_undeclared_variable.mc -o bad_name.o --emit-ir bad_name.ll",
        "Два сообщения semantic error с указанием строки и столбца.",
    )
    cursor = add_placeholder_after(
        cursor,
        "Short-circuit для логического &&",
        f"cd \"{WORKTREE_WSL}/build\"\n./mini_cc ../tests/10_short_circuit.mc -o sc.o --emit-ir sc.ll\nclang ../runtime/main.c sc.o -o sc_app\n./sc_app",
        "Вывод 123 без аварийного деления на ноль.",
    )
    cursor = add_placeholder_after(
        cursor,
        "Прогон теста суммы простых чисел",
        f"cd \"{WORKTREE_WSL}/build\"\n./mini_cc ../tests/12_sum_primes.mc -o primes.o --emit-ir primes.ll\nclang ../runtime/main.c primes.o -o primes_app\n./primes_app",
        "Вывод 17 для compiled_fn(10).",
    )
    cursor = add_placeholder_after(
        cursor,
        "Полный прогон положительных тестов",
        f"cd \"{WORKTREE_WSL}/build\"\nfor t in ../tests/01_return_literal.mc ../tests/02_arithmetic.mc ../tests/03_if_else.mc ../tests/04_for.mc ../tests/05_ternary_true.mc ../tests/06_ternary_false.mc ../tests/07_nested_ternary.mc ../tests/10_short_circuit.mc ../tests/11_function_call.mc ../tests/12_sum_primes.mc ../tests/test.mc; do\n  base=$(basename \"$t\" .mc)\n  ./mini_cc \"$t\" -o \"$base.o\" --emit-ir \"$base.ll\"\n  clang ../runtime/main.c \"$base.o\" -o \"$base.app\"\n  printf \"%s: \" \"$base\"\n  ./$base.app\ndone",
        "Последовательный вывод результатов всех положительных тестов.",
    )


def add_appendices(doc: Document) -> None:
    if has_heading(doc, "13. Приложения"):
        return

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading = doc.add_heading("13. Приложения", level=1)
    format_heading(heading, 1)

    intro = doc.add_paragraph(
        "В приложениях приведены полные листинги основных файлов проекта. "
        "Это увеличивает воспроизводимость отчёта и позволяет показать, "
        "что реализованный компилятор действительно состоит из самостоятельных модулей front-end, "
        "семантического анализа и генерации кода."
    )
    format_body(intro)

    files = [
        ("Приложение А. Полный листинг src/lexer.l", "src/lexer.l"),
        ("Приложение Б. Полный листинг src/parser.y", "src/parser.y"),
        ("Приложение В. Полный листинг src/ast.hpp", "src/ast.hpp"),
        ("Приложение Г. Полный листинг src/ast.cpp", "src/ast.cpp"),
        ("Приложение Д. Полный листинг src/sema.hpp", "src/sema.hpp"),
        ("Приложение Е. Полный листинг src/sema.cpp", "src/sema.cpp"),
        ("Приложение Ж. Полный листинг src/codegen.hpp", "src/codegen.hpp"),
        ("Приложение З. Полный листинг src/codegen.cpp", "src/codegen.cpp"),
        ("Приложение И. Полный листинг src/driver.cpp", "src/driver.cpp"),
        ("Приложение К. Полный листинг runtime/main.c", "runtime/main.c"),
        ("Приложение Л. Полный листинг CMakeLists.txt", "CMakeLists.txt"),
        ("Приложение М. Полный листинг scripts/run_tests.sh", "scripts/run_tests.sh"),
        ("Приложение Н. Пример содержательного теста tests/12_sum_primes.mc", "tests/12_sum_primes.mc"),
    ]

    for title, rel_path in files:
        h = doc.add_heading(title, level=2)
        format_heading(h, 2)
        desc = doc.add_paragraph(
            f"Ниже приведён полный текст файла {rel_path}. "
            "Листинг оставлен без сокращений, чтобы при необходимости можно было "
            "сверить его с репозиторием и повторить сборку проекта."
        )
        format_body(desc)

        code_p = doc.add_paragraph()
        code_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        code_p.paragraph_format.first_line_indent = Cm(0)
        code_p.paragraph_format.left_indent = Cm(0.5)
        code_p.paragraph_format.space_after = Pt(6)
        run = code_p.add_run(read_text(rel_path).rstrip())
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8)


def update_guide() -> None:
    GUIDE.write_text(
        f"""# Как сделать скриншоты для отчёта

Все команды лучше выполнять в **WSL Ubuntu**. На каждом скриншоте желательно оставлять видимыми и команду, и её результат. Ниже перечислены места, которые добавлены в отчёт, и команды для их заполнения.

## 1. Структура проекта

```bash
cd "{WORKTREE_WSL}"
ls
find src runtime tests -maxdepth 1 -type f | sort
```

Вставить в раздел 10 рядом с местом `Структура проекта`.

## 2. Успешная сборка компилятора

```bash
cd "{WORKTREE_WSL}"
mkdir -p build
cd build
cmake ..
cmake --build .
```

Вставить в раздел 10 рядом с местом `Успешная сборка компилятора`.

## 3. Запуск основной тестовой программы

```bash
cd "{WORKTREE_WSL}/build"
./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll
clang ../runtime/main.c generated.o -o app
./app
```

На скриншоте должен быть виден вывод:

```text
14
```

## 4. Просмотр PHI-узла в LLVM IR

```bash
cd "{WORKTREE_WSL}/build"
./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll
grep -n phi generated.ll
```

Вставить в раздел 7 рядом с местом `Фрагмент LLVM IR с PHI-узлом`.

## 5. Семантические ошибки

```bash
cd "{WORKTREE_WSL}/build"
./mini_cc ../tests/08_type_error.mc -o bad_type.o --emit-ir bad_type.ll
./mini_cc ../tests/09_undeclared_variable.mc -o bad_name.o --emit-ir bad_name.ll
```

Ожидается два сообщения `semantic error`.

## 6. Проверка short-circuit

```bash
cd "{WORKTREE_WSL}/build"
./mini_cc ../tests/10_short_circuit.mc -o sc.o --emit-ir sc.ll
clang ../runtime/main.c sc.o -o sc_app
./sc_app
```

Ожидаемый вывод:

```text
123
```

## 7. Проверка теста суммы простых чисел

```bash
cd "{WORKTREE_WSL}/build"
./mini_cc ../tests/12_sum_primes.mc -o primes.o --emit-ir primes.ll
clang ../runtime/main.c primes.o -o primes_app
./primes_app
```

Ожидаемый вывод:

```text
17
```

## 8. Полный прогон положительных тестов

```bash
cd "{WORKTREE_WSL}/build"
for t in ../tests/01_return_literal.mc ../tests/02_arithmetic.mc ../tests/03_if_else.mc ../tests/04_for.mc ../tests/05_ternary_true.mc ../tests/06_ternary_false.mc ../tests/07_nested_ternary.mc ../tests/10_short_circuit.mc ../tests/11_function_call.mc ../tests/12_sum_primes.mc ../tests/test.mc; do
  base=$(basename "$t" .mc)
  ./mini_cc "$t" -o "$base.o" --emit-ir "$base.ll"
  clang ../runtime/main.c "$base.o" -o "$base.app"
  printf "%s: " "$base"
  ./$base.app
done
```

Ожидаемый вывод:

```text
01_return_literal: 42
02_arithmetic: 35
03_if_else: 11
04_for: 45
05_ternary_true: 100
06_ternary_false: 200
07_nested_ternary: 1
10_short_circuit: 123
11_function_call: 15
12_sum_primes: 17
test: 14
```

## 9. Как аккуратно вставлять скриншоты в Word

1. Делай скриншот так, чтобы были видны и команда, и результат.
2. Обрезай только лишние пустые поля, не удаляя имя файла, путь и сам вывод.
3. В Word вставляй изображение прямо в таблицу-заглушку или сразу под неё.
4. Если скриншот получается высоким, уменьшай ширину, но не делай текст нечитаемым.
5. После вставки можно заменить строку `МЕСТО ДЛЯ СКРИНШОТА: ...` на подпись к рисунку.
""",
        encoding="utf-8",
    )


def main() -> None:
    doc = Document(REPORT)
    add_extended_sections(doc)
    add_appendices(doc)
    try:
        doc.save(REPORT)
        saved_to = REPORT
    except PermissionError:
        saved_to = REPORT.with_name(REPORT.stem + "_расширенный.docx")
        doc.save(saved_to)
    update_guide()
    print(saved_to)
    print(GUIDE)


if __name__ == "__main__":
    main()
