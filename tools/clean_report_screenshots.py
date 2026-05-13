from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Отчёт_мини_компилятор_вариант_11_с_простыми.docx"
OUT = ROOT / "Отчёт_мини_компилятор_вариант_11_проверенный.docx"


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def has_image(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def main() -> None:
    doc = Document(SRC)
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "compiled_fn(20):":
            remove_paragraph(paragraph)
            if index + 1 < len(paragraphs) and has_image(paragraphs[index + 1]):
                remove_paragraph(paragraphs[index + 1])
            break
    doc.save(OUT)

    check = Document(OUT)
    all_text = "\n".join(p.text for p in check.paragraphs)
    print(OUT)
    print("inline_shapes", len(check.inline_shapes))
    print("has_compiled_fn_20", "compiled_fn(20)" in all_text)


if __name__ == "__main__":
    main()
